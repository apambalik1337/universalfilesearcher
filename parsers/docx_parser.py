"""
docx_parser.py

Parses modern Word documents (.docx) using python-docx.

Requirements from spec (Priority 4 - MUST WORK):
  - search paragraphs
  - search tables / table cells (not just normal paragraphs)
  - search headings where available
  - a name inside a table must still be found
"""

import docx

from .base_parser import BaseParser, ParseResult, Record


class DocxParser(BaseParser):
    extensions = (".docx",)
    display_name = "Word (.docx)"

    def parse(self, file_path: str) -> ParseResult:
        try:
            document = docx.Document(file_path)
        except Exception as e:
            return ParseResult(success=False, error=f"Failed to open document: {e}")

        records = []

        # --- Paragraphs (includes headings, since headings are just
        # paragraphs with a "Heading" style in python-docx) ---
        for i, para in enumerate(document.paragraphs, start=1):
            text = (para.text or "").strip()
            if not text:
                continue
            is_heading = para.style is not None and para.style.name and "Heading" in para.style.name
            location = f"Heading (Paragraph {i})" if is_heading else f"Paragraph {i}"
            records.append(Record(content=text, location=location))

        # --- Tables ---
        for t_idx, table in enumerate(document.tables, start=1):
            for r_idx, row in enumerate(table.rows, start=1):
                for c_idx, cell in enumerate(row.cells, start=1):
                    text = (cell.text or "").strip()
                    if not text:
                        continue
                    col_name = _table_column_header(table, c_idx)
                    records.append(
                        Record(
                            content=text,
                            location=f"Table {t_idx}",
                            row=r_idx,
                            column=col_name,
                        )
                    )

        if not records:
            return ParseResult(success=True, records=[], warning="No content found in document")

        return ParseResult(success=True, records=records)

    def is_dependency_available(self) -> bool:
        try:
            import docx  # noqa
            return True
        except ImportError:
            return False

    def dependency_message(self) -> str:
        return "Install with: pip install python-docx"


def _table_column_header(table, col_idx: int):
    """Best-effort: use the first row as a header label for this column, if present."""
    try:
        header_row = table.rows[0]
        if col_idx - 1 < len(header_row.cells):
            text = header_row.cells[col_idx - 1].text.strip()
            return text or None
    except Exception:
        pass
    return None
